##########################################
# Project name  : Outlook Support Classification Agent — ask_user_params generator (one-time / local tool)
# Date          : 2026-08-03
# purpose       : Generate ask_user_params.json entries with an LLM. For each KB it pairs the KB's article (.docx in
#                 ./kb_articles) with its PowerShell script (.ps1/.txt in ./scripts), both named by the KB number, and
#                 sends them to GPT-5 mini in one STATELESS, structured call. It writes an entry in the exact
#                 {script, params, allowed_values, examples} shape. allowed_values are only a FEW FORMAT EXAMPLES per
#                 parameter (never an enforced/complete list) — so it works generically for any value set (time zones,
#                 add-ins, folders, ...), which may have hundreds of valid values that cannot be enumerated.
# Connections   : Reads settings from config.yaml; reads <KB>.docx + <KB>.ps1|.txt; writes ask_user_params.json (copy it
#                 into the validation function's folder to deploy). Auth: Entra ID via `az login` (no keys).
##########################################

# =====================================================================
# IMPORTS
# =====================================================================
from __future__ import annotations  # Enable postponed evaluation of type annotations (PEP 563)

import argparse  # Parse command-line options
import json  # Parse the LLM output + read/write ask_user_params.json
import sys  # Exit with a clear status on setup problems
from pathlib import Path  # Filesystem paths
from typing import Any, Optional  # Type hints

import yaml  # Load the tool's config.yaml
from azure.identity import DefaultAzureCredential, get_bearer_token_provider  # Entra ID auth (no keys)
from openai import AzureOpenAI  # Azure OpenAI SDK client (GPT-5 mini)

try:  # python-docx extracts the article text from the .docx
    import docx  # python-docx
except ImportError:  # Fail with guidance if it isn't installed
    docx = None  # Sentinel checked in main()

_SCRIPT_EXTENSIONS = (".ps1", ".txt")  # Accepted PowerShell script file extensions (first match wins)

# System prompt: turns (PowerShell script + KB article) into the ask_user_params inner object.
_SYSTEM_PROMPT = (
    "You generate a JSON parameter specification for an Outlook self-help automation. You are given a PowerShell "
    "remediation script and the text of the KB article it automates. Produce the JSON an assistant will use to collect "
    "the script's inputs from users who type in natural language.\n\n"
    "From the PowerShell param() block, extract each parameter:\n"
    "- name: the parameter name (without the $).\n"
    "- required: true if the parameter is Mandatory OR has no default value; false if it has a default.\n"
    "- description: a short, user-facing explanation (use the KB article for context).\n\n"
    "allowed_values: for each parameter give a FEW representative EXAMPLE values that show the expected FORMAT. This is "
    "GUIDANCE ONLY, never an exhaustive list — do not try to enumerate every possible value (a set may have hundreds, "
    "e.g. time zones or add-ins). Even when the script has a [ValidateSet(...)], include just a few of its values as "
    "examples (you need not list them all).\n\n"
    "examples: 2 realistic natural-language requests a user might type, each mapped to the parameter values.\n\n"
    "RULES:\n"
    "- Treat allowed_values purely as format examples — do NOT imply it is the complete set.\n"
    "- Do NOT include a 'script' field — the tool sets it from the filename.\n"
    "- Output ONLY a JSON object, no prose, no code fences, in EXACTLY this shape:\n"
    "{\"params\": [{\"name\": <string>, \"required\": <bool>, \"description\": <string>}], "
    "\"allowed_values\": {<param name>: [<example value strings>]}, "
    "\"examples\": [{\"prompt\": <string>, \"params\": {<param name>: <value>}}]}"
)


# =====================================================================
# HELPERS
# =====================================================================
def _extract_first_json_object(text: str) -> Optional[dict[str, Any]]:  # Pull the first balanced {...} from model text
    """Return the first balanced JSON object in the text (tolerant of prose/code fences), or None."""
    start = text.find("{")  # Locate the first opening brace
    if start == -1:  # No object present
        return None  # Nothing to parse
    depth = 0  # Current brace nesting depth
    in_string = False  # Inside a JSON string literal?
    escaped = False  # Previous char was an escape?
    for position in range(start, len(text)):  # Scan from the first brace
        character = text[position]  # Current character
        if in_string:  # Inside a string
            if escaped:  # Consume an escaped char
                escaped = False
            elif character == "\\":  # Start of an escape
                escaped = True
            elif character == '"':  # End of the string
                in_string = False
            continue  # Keep scanning
        if character == '"':  # String opens
            in_string = True
        elif character == "{":  # Nested object opens
            depth += 1
        elif character == "}":  # Object closes
            depth -= 1
            if depth == 0:  # Balanced back to the top-level object
                try:  # Parse the balanced span
                    parsed = json.loads(text[start : position + 1])
                except ValueError:  # Not valid JSON
                    return None
                return parsed if isinstance(parsed, dict) else None  # Only return objects
    return None  # Never balanced


def _read_docx_text(docx_path: Path) -> str:  # Extract paragraph + table text from a .docx
    """Extract readable text (paragraphs + tables) from a .docx via python-docx. Images are ignored (not needed here)."""
    document = docx.Document(str(docx_path))  # Open the .docx
    lines: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]  # Non-empty paragraphs
    for table in document.tables:  # Include table cell text
        for row in table.rows:  # Each row
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]  # Non-empty cells
            if cells:  # Only rows with content
                lines.append(" | ".join(cells))  # Join the row's cells
    return "\n".join(lines)  # The article text


def _find_pairs(kb_articles_dir: Path, scripts_dir: Path, only_kb: Optional[str]) -> list[tuple[str, Path, Path]]:
    """Find (kb_id, docx_path, script_path) triples: each <kb>.docx in kb_articles_dir with a <kb>.ps1/.txt in scripts_dir."""
    triples: list[tuple[str, Path, Path]] = []  # Collected pairs
    for docx_path in sorted(kb_articles_dir.glob("*.docx")):  # Every article
        if docx_path.name.startswith("~$"):  # Skip Word's temporary lock files
            continue
        kb_id = docx_path.stem  # KB id = the filename without extension (e.g. KB001234)
        if only_kb and kb_id != only_kb:  # Filter to one KB when --kb is given
            continue
        script_path = next(  # First matching script in the scripts folder (<kb>.ps1 then <kb>.txt)
            (scripts_dir / f"{kb_id}{ext}" for ext in _SCRIPT_EXTENSIONS if (scripts_dir / f"{kb_id}{ext}").is_file()),
            None,
        )
        if script_path is None:  # No script for this KB -> skip with a note
            print(f"  ! {kb_id}: no {kb_id}.ps1/.txt found in {scripts_dir} — skipping")
            continue
        triples.append((kb_id, docx_path, script_path))  # A complete pair
    return triples  # All pairs found


def _sanity_check(kb_id: str, entry: dict[str, Any]) -> None:  # Print non-fatal warnings about the generated entry
    """Warn (do not fail) when the generated entry looks off, so you can review before deploying."""
    declared = {p.get("name") for p in entry.get("params", []) or []}  # Declared parameter names
    allowed = entry.get("allowed_values", {}) or {}  # Per-param example values (format guidance)
    for param in entry.get("params", []) or []:  # Each declared param
        if not allowed.get(param.get("name")):  # No example values -> the model has no format guidance for it
            print(f"    ⚠ {kb_id}: param {param.get('name')!r} has no example values — consider adding a couple")
    for example in entry.get("examples", []) or []:  # Each example
        for name in (example.get("params", {}) or {}):  # Each param used in the example
            if name not in declared:  # References a param not in the params list
                print(f"    ⚠ {kb_id}: example references undeclared param {name!r} — review")


def _generate_entry(client: AzureOpenAI, deployment: str, docx_path: Path, script_path: Path) -> dict[str, Any]:  # One KB
    """Read the article + script, call GPT-5 mini once (stateless), and assemble the ask_user_params entry."""
    article_text = _read_docx_text(docx_path)  # KB article text (images ignored)
    script_text = script_path.read_text(encoding="utf-8", errors="replace")  # PowerShell script text
    script_name = script_path.name if script_path.suffix.lower() == ".ps1" else f"{script_path.stem}.ps1"  # Bare .ps1 name
    user_content = (  # The single stateless payload: script + article
        f"POWERSHELL SCRIPT (filename: {script_name}):\n```powershell\n{script_text}\n```\n\n"
        f"KB ARTICLE TEXT:\n{article_text}"
    )
    response = client.chat.completions.create(  # Stateless chat call (NO temperature — GPT-5 family)
        model=deployment,  # The GPT-5 mini deployment
        messages=[  # Fixed instruction + this KB's payload
            {"role": "system", "content": _SYSTEM_PROMPT},  # How to produce the JSON
            {"role": "user", "content": user_content},  # The script + article
        ],
    )
    raw_text = response.choices[0].message.content or ""  # The model's text output
    inner = _extract_first_json_object(raw_text)  # Parse the JSON object
    if inner is None:  # The model didn't return JSON
        raise ValueError(f"LLM did not return a JSON object. First 500 chars:\n{raw_text[:500]}")
    return {  # Assemble the final entry; the tool (not the model) sets the script filename
        "script": script_name,  # The exact .ps1 the executor will run (bare filename)
        "params": inner.get("params", []),  # Parameter definitions (name / required / description)
        "allowed_values": inner.get("allowed_values", {}),  # Per-param FORMAT EXAMPLES (not an enforced/complete list)
        "examples": inner.get("examples", []),  # Natural-language -> params examples
    }


def _load_config(config_path: Path) -> dict[str, Any]:  # Load the tool's config.yaml (with sensible fallbacks)
    """Load config.yaml; return a dict with azure_openai + paths sections (empty defaults if missing)."""
    if not config_path.is_file():  # Config is optional; CLI flags can supply everything
        return {"azure_openai": {}, "paths": {}}
    data = yaml.safe_load(config_path.read_text(encoding="utf-8")) or {}  # Parse YAML
    data.setdefault("azure_openai", {})  # Ensure sections exist
    data.setdefault("paths", {})
    return data


# =====================================================================
# ENTRY POINT
# =====================================================================
def main() -> int:  # Command-line entry point
    """Generate ask_user_params.json entries for the KB .docx/.ps1 pairs found under ./kb_articles + ./scripts."""
    here = Path(__file__).resolve().parent  # This tool's folder
    parser = argparse.ArgumentParser(description="Generate ask_user_params.json entries with GPT-5 mini.")  # CLI
    parser.add_argument("--config", default=str(here / "config.yaml"), help="Path to config.yaml")
    parser.add_argument("--kb", default=None, help="Only process this KB id (default: all pairs found)")
    parser.add_argument("--deployment", default=None, help="GPT-5 mini deployment name (overrides config.yaml)")
    parser.add_argument("--endpoint", default=None, help="Azure OpenAI endpoint (overrides config.yaml)")
    args = parser.parse_args()  # Parse CLI

    if docx is None:  # python-docx is required to read the articles
        print("ERROR: python-docx is not installed. Run: pip install -r requirements.txt", file=sys.stderr)
        return 2

    config = _load_config(Path(args.config))  # Load config.yaml
    aoai = config["azure_openai"]  # Azure OpenAI settings
    paths = config["paths"]  # Folder settings
    endpoint = args.endpoint or aoai.get("endpoint", "")  # Endpoint (CLI > config)
    deployment = args.deployment or aoai.get("deployment", "")  # Deployment (CLI > config)
    api_version = aoai.get("api_version", "2024-12-01-preview")  # API version
    token_scope = aoai.get("token_scope", "https://ai.azure.com/.default")  # Entra token scope

    if not endpoint or not deployment or deployment.startswith("<"):  # Must be configured
        print("ERROR: set azure_openai.endpoint and azure_openai.deployment in config.yaml (or via --endpoint/--deployment).",
              file=sys.stderr)
        return 2

    kb_articles_dir = here / paths.get("kb_articles_dir", "kb_articles")  # Article folder
    scripts_dir = here / paths.get("scripts_dir", "scripts")  # Script folder
    output_path = here / paths.get("output", "ask_user_params.json")  # Output file
    for folder in (kb_articles_dir, scripts_dir):  # Both input folders must exist
        if not folder.is_dir():
            print(f"ERROR: input folder not found: {folder} (create it and add your files).", file=sys.stderr)
            return 2

    pairs = _find_pairs(kb_articles_dir, scripts_dir, args.kb)  # Find the KB .docx/.ps1 pairs
    if not pairs:  # Nothing to do
        print(f"No <KB>.docx (in {kb_articles_dir.name}) with a matching <KB>.ps1/.txt (in {scripts_dir.name}) found"
              + (f" for kb {args.kb!r}." if args.kb else "."))
        return 1

    # Build the Azure OpenAI client using Entra ID (run `az login` first; no keys).
    credential = DefaultAzureCredential()  # Uses your az-login / VS Code Azure identity locally
    token_provider = get_bearer_token_provider(credential, token_scope)  # Bearer tokens for the endpoint
    client = AzureOpenAI(azure_endpoint=endpoint, api_version=api_version, azure_ad_token_provider=token_provider)  # Client

    # Load any existing output so we MERGE (running per-KB accumulates entries).
    data: dict[str, Any] = {}  # The merged spec map
    if output_path.is_file():  # Reuse existing entries
        try:
            data = json.loads(output_path.read_text(encoding="utf-8")) or {}  # Load current file
        except ValueError:
            print(f"WARNING: {output_path.name} was not valid JSON; starting fresh.")

    generated = 0  # Count of successful entries
    for kb_id, docx_path, script_path in pairs:  # Process each pair
        print(f"Generating {kb_id}  ({docx_path.name} + {script_path.name}) ...")
        try:  # Generate one entry via the LLM
            entry = _generate_entry(client, deployment, docx_path, script_path)
        except Exception as generation_error:  # Never let one KB abort the whole run
            print(f"  ! {kb_id}: failed: {generation_error}")
            continue
        _sanity_check(kb_id, entry)  # Print any review warnings
        data[kb_id] = entry  # Merge into the map
        generated += 1  # Tally
        print(f"  ✓ {kb_id}: script={entry['script']}, {len(entry['params'])} param(s), "
              f"{len(entry['examples'])} example(s)")

    output_path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")  # Write the merged file
    print(f"\nWrote {generated} entr(y/ies) to {output_path}")  # Summary
    print("Review the JSON (allowed_values are format examples, not enforced lists), then copy it into "
          "the validation function's folder to deploy.")
    return 0  # Success


if __name__ == "__main__":  # Allow `python generate_ask_user_params.py`
    raise SystemExit(main())  # Exit with the return code
